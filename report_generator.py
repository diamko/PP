# report_generator.py
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Установка цвета заднего фона ячейки таблицы"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Установка внутренних отступов (Padding) в ячейке для красоты"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_order_document(path, astronaut_id, fio, gender, suit_mod, anthro_params, calculated_data):
    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(18)
    run_title = title_p.add_run("ТЕХНИЧЕСКИЙ ЗАКАЗ-НАРЯД\nНА ПОДБОР И ИЗГОТОВЛЕНИЕ СНАРЯЖЕНИЯ СК «СОКОЛ-КВ»")
    run_title.bold = True
    run_title.size = Pt(14)
    run_title.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    p_info = doc.add_paragraph()
    p_info.paragraph_format.space_after = Pt(12)
    p_info.paragraph_format.line_spacing = 1.25

    p_info.add_run("Личное дело космонавта: ").bold = True
    p_info.add_run(f"{fio} ({gender[0].upper() if gender else 'М'})\n")
    p_info.add_run("ID Личного дела: ").bold = True
    p_info.add_run(f"{astronaut_id}   |   ")
    p_info.add_run("Модификация скафандра: ").bold = True
    run_mod = p_info.add_run(f"{suit_mod}\n")
    run_mod.bold = True

    p_info.add_run("Исходная антропометрия: ").bold = True
    p_info.add_run(
        f"Голова: {anthro_params.get('head', '-')} см, "
        f"Рост: {anthro_params.get('height', '-')} см, "
        f"Грудь: {anthro_params.get('chest', '-')} см, "
        f"Талия: {anthro_params.get('waist', '-')} см, "
        f"Обувь: {anthro_params.get('shoe', '-')} разм., "
        f"Обхват запястья: {anthro_params.get('wrist_circ', '-')} см, "
        f"Длина пальца: {anthro_params.get('finger_len', '-')} см, "
        f"Длина руки: {anthro_params.get('arm_len', '-')} см, "
        f"Длина ноги: {anthro_params.get('leg_len', '-')} см."
    )

    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.add_run("Расчетные параметры изготавливаемого комплекта изделия:").bold = True

    table = doc.add_table(rows=2, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = [
        "№ Изд.", "Именной\nиндекс", "Размер\nоболочки", "ГП-7С",
        "ШЛ-10СА", "Белье", "Носки\n(ГОСТ)", "Стельки", "Перчатки\nстарт.", "Обувь\nстарт."
    ]

    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(hdr_cells[i], "F2F2F2")
        set_cell_margins(hdr_cells[i], top=80, bottom=80, left=60, right=60)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

    product_id = str(calculated_data.get("product_id", astronaut_id + 500))

    data_fields = [
        product_id,
        calculated_data.get("name_index", "КС"),
        calculated_data.get("suit_size", "-"),
        calculated_data.get("gloves_gp", "-"),
        calculated_data.get("shl_size", "-"),
        calculated_data.get("underwear_size", "-"),
        calculated_data.get("socks_size", "-"),
        calculated_data.get("insoles_size", "-"),
        calculated_data.get("gloves_start", "-"),
        calculated_data.get("boots_size", "-")
    ]

    data_cells = table.rows[1].cells
    for i, field_value in enumerate(data_fields):
        data_cells[i].text = str(field_value)
        data_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        data_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(data_cells[i], top=100, bottom=100, left=60, right=60)
        data_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(24)
    h3.paragraph_format.space_after = Pt(6)
    h3.add_run("Сопутствующее санитарно-гигиеническое обеспечение:").bold = True

    p_socks = doc.add_paragraph(style='List Bullet')
    p_socks.paragraph_format.space_after = Pt(3)
    p_socks.add_run("Трусы эластичные под скафандр (Seni Active): ").bold = True
    p_socks.add_run(f"Типоразмер {calculated_data.get('seni_size', 'M')}")

    p_gloves = doc.add_paragraph(style='List Bullet')
    p_gloves.paragraph_format.space_after = Pt(12)
    p_gloves.add_run("Перчатки кольчужные хирургические (защитные): ").bold = True
    p_gloves.add_run(f"Типоразмер {calculated_data.get('gloves_chainmail', 'M')}")

    p_note = doc.add_paragraph()
    p_note.paragraph_format.left_indent = Inches(0.2)
    p_note.paragraph_format.space_before = Pt(6)
    run_note_lbl = p_note.add_run("ВАЖНОЕ ПРИМЕЧАНИЕ ПО ЭКСПЛУАТАЦИИ: ")
    run_note_lbl.bold = True
    run_note_lbl.font.size = Pt(10)

    run_note_txt = p_note.add_run(
        "Под герметичную полетную перчатку ГП-7С на руку оператора в обязательном порядке надевается "
        "сначала хирургическая кольчужная перчатка (для защиты кожных покровов), а поверх неё — "
        "стартовая технологическая перчатка для обеспечения плотной фиксации ложемента кисти."
    )
    run_note_txt.font.italic = True
    run_note_txt.font.size = Pt(10)

    doc.save(path)